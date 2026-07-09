"""Microsoft Word document parser."""

from __future__ import annotations

from pathlib import Path

from custom_rag.core.exceptions import ParseFailedError
from custom_rag.core.types import BlockType, DocumentMetadata, ParsedDocument
from custom_rag.ingestion.parsers.base import BaseParser, BlockBuilder

_HEADING_STYLES = {
    "Title": 1,
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
}


class DOCXParser(BaseParser):
    name = "docx"
    supported_extensions = frozenset({".docx"})
    supported_mimes = frozenset(
        {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    )

    def parse(self, path: Path, metadata: DocumentMetadata) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise ParseFailedError(
                str(path),
                "python-docx is not installed; use pip install 'custom-rag[ingestion]'",
                cause=exc,
            ) from exc

        try:
            document = Document(str(path))
        except Exception as exc:
            raise ParseFailedError(str(path), "unable to open DOCX file", cause=exc) from exc

        builder = BlockBuilder()
        section_stack: list[tuple[str, int]] = []
        section_counter = 0
        paragraph_counter = 0
        row_counter = 0
        raw_parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name if paragraph.style else ""
            if style_name in _HEADING_STYLES:
                section_counter += 1
                block_id = f"sec_{section_counter}"
                level = _HEADING_STYLES[style_name]
                while section_stack and section_stack[-1][1] >= level:
                    section_stack.pop()
                parent_id = section_stack[-1][0] if section_stack else None
                hierarchy = ["doc", *[item[0] for item in section_stack], block_id]
                builder.add(
                    block_id=block_id,
                    block_type=BlockType.SECTION,
                    text=text,
                    parent_block_id=parent_id,
                    order=section_counter,
                    hierarchy_path=hierarchy,
                    metadata={"heading_level": level, "style": style_name},
                )
                section_stack.append((block_id, level))
                raw_parts.append(text)
                continue

            paragraph_counter += 1
            parent_id, hierarchy = _current_parent(section_stack)
            block_id = f"p_{paragraph_counter}"
            builder.add(
                block_id=block_id,
                block_type=BlockType.PARAGRAPH,
                text=text,
                parent_block_id=parent_id,
                order=paragraph_counter,
                hierarchy_path=[*hierarchy, block_id],
                metadata={"style": style_name or "Normal"},
            )
            raw_parts.append(text)

        for table in document.tables:
            parent_id, hierarchy = _current_parent(section_stack)
            for row_index, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if not cells:
                    continue
                row_counter += 1
                row_text = " | ".join(cells)
                block_id = f"row_{row_counter}"
                builder.add(
                    block_id=block_id,
                    block_type=BlockType.TABLE_ROW,
                    text=row_text,
                    parent_block_id=parent_id,
                    order=row_counter,
                    hierarchy_path=[*hierarchy, block_id],
                    metadata={"row_index": row_index},
                )
                raw_parts.append(row_text)

        doc_metadata = metadata.model_copy(update={"doc_type": "docx"})
        return ParsedDocument(
            metadata=doc_metadata,
            blocks=builder.blocks,
            raw_text="\n\n".join(raw_parts) if raw_parts else None,
        )


def _current_parent(section_stack: list[tuple[str, int]]) -> tuple[str | None, list[str]]:
    if not section_stack:
        return None, ["doc"]
    hierarchy = ["doc", *[item[0] for item in section_stack]]
    return section_stack[-1][0], hierarchy
